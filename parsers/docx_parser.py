from __future__ import annotations

import os

from docx import Document

from .base import ParsedDocument, ParsedPage
from utils.logger import get_logger

logger = get_logger(__name__)


def parse_docx(file_path: str) -> ParsedDocument:
    """Extract text and tables from a DOCX file."""
    doc = Document(file_path)
    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            # Preserve heading structure
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                try:
                    lvl = int(level)
                except ValueError:
                    lvl = 1
                prefix = "#" * lvl
                text_parts.append(f"{prefix} {line}")
            else:
                text_parts.append(line)

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    # Convert tables to text
    for tbl in tables:
        if len(tbl) >= 2:
            header = tbl[0]
            text_parts.append(f"\n[表格] {' | '.join(header)}")
            for row in tbl[1:]:
                kv = [f"{h}: {v}" for h, v in zip(header, row) if h.strip() and v.strip()]
                if kv:
                    text_parts.append("; ".join(kv))

    page = ParsedPage(
        page_number=1,
        text="\n".join(text_parts),
        tables=tables,
    )

    logger.info(f"DOCX parsed: {file_path} ({len(doc.paragraphs)} paragraphs, {len(tables)} tables)")
    return ParsedDocument(
        file_name=os.path.basename(file_path),
        file_path=file_path,
        file_type="docx",
        pages=[page],
        metadata={"total_paragraphs": len(doc.paragraphs), "total_tables": len(tables)},
    )
